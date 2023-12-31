#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Proliverick is a small app writtten in Python which generates excersises for foreing language students 
according to particular patterns using the OpenAI API.
@copyright: Copyright © 2023 Iwan van der Kleijn
@license: MIT
"""
import time
from typing import List
from enum import Enum
from pathlib import Path
import pandas as pd
from docx import Document
from simpleaichat import AIChat
from dotenv import load_dotenv

from proliverick import prompts
        
class OutputType(Enum):
    TEXT = 'Text'
    WORD = 'Word document'
    EXCEL = 'Excel sheet'
    
    @staticmethod
    def initialize_from_str(value_str):
        for output_type in OutputType:
            if output_type.value == value_str:
                return output_type
        raise ValueError(f"'{value_str}' is not a valid OutputType value")


class AssesementCreator:
    
    def __init__(self):
        load_dotenv()
        #params = {"temperature": 1.5}  #"max_tokens": 100}  # a temperature of 0.0 is deterministic
        #self.ai = AIChat(model="gpt-4", console= False, params=params)  #, system='You are a system which generated questions for foreign languages students' )

        self.ai = AIChat(model="gpt-4", console= False)
        
    def get_inverse_interrogative(self, num: int) -> List[str]:
        prompt  = prompts.inverse_interrogative_prompt(num)
        return self.ai(prompt)    
     
    @staticmethod
    def write_to_file(output_type: OutputType, output_path: Path, title: str, data: str):
        
        data_list = data.split('\n')
        match output_type:
            case OutputType.TEXT:
                with output_path.open('w') as f:
                    f.write(title)
                    f.write('\n')   
                    f.write(data)
            case OutputType.WORD:
                doc = Document()
                doc.add_heading(title, 0)
                for line in data_list:
                    doc.add_paragraph(line)
                doc.save(output_path)
            case OutputType.EXCEL:
                #data_list.insert(0, title)
                #df = pd.DataFrame(data_list, columns=["Data"])
                df = pd.DataFrame(data_list, columns=[title])
                df.to_excel(output_path, index=False)
            case _:
                raise ValueError(f"Unsupported output type: {output_type}")


def create_assesement(data: dict):
    creator = AssesementCreator()
    if data['element-type'] == "Reverse interrogation":
        #items = "lala\nlala"
        items = creator.get_inverse_interrogative(data['number-of-elements'])
        #time.sleep(5)
        output_type = OutputType.initialize_from_str(data['output-file-type'])
        output_path = Path(data['output-file-path'])
        title = data['title']
        creator.write_to_file(output_type, output_path, title, items)
    else:
        raise ValueError(f"Unsupported element type: {data['element-type']}")
        #AssesementCreator.write_to_file(data['output-file-type'], Path(data['output-file-path']), data)

